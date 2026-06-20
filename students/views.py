# students/views.py - منظم مع الصلاحيات

import csv
import io
import json
from datetime import date, datetime, timedelta
from decimal import Decimal

import openpyxl
from django.contrib import messages
from django.core.paginator import Paginator
from django.db.models import Q, Sum, Count, Avg, Max, Min
from django.http import JsonResponse, HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse
from django.utils import timezone
from django.views.decorators.cache import never_cache
from openpyxl.styles import Font, PatternFill, Alignment

from .decorators import (
    students_basic_access,
    students_add_only,
    students_edit_access,
    students_archive_access,
    students_import_export_access,
    students_full_access,
    students_reports_access,
    students_sensitive_operation,
)

from .forms import StudentForm, Student_edit_Form

from .models import (
    Student,
    ArchiveStudent,
    UserProfile,
)

from .services.export_service import StudentExportService
from .services.financial_service import StudentFinancialService
from .services.import_preview_service import StudentImportPreviewService

from school_settings.models import (
    AcademicYear as SettingsAcademicYear,
    EducationLevel,
    GradeLevel,
    SchoolFeesSettings,
    SystemSettings,
)


try:
    from .utils.upgrade_utils import StudentUpgradeManager
    UPGRADE_AVAILABLE = True
except ImportError:
    UPGRADE_AVAILABLE = False
    StudentUpgradeManager = None

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ===================================
# 🔧 الدوال المساعدة
# ===================================

def get_user_role(user):
    """الحصول على دور المستخدم"""
    if hasattr(user, 'system_role'):
        return user.system_role.role
    return None

def get_system_data():
    """الحصول على بيانات النظام الأساسية"""
    try:
        system_settings = SystemSettings.get_current_settings()
    except:
        system_settings = None
        
    try:
        current_year = SettingsAcademicYear.get_current_year()
    except:
        current_year = None
        
    return system_settings, current_year

def get_education_data():
    """الحصول على البيانات التعليمية"""
    education_levels = EducationLevel.objects.filter(is_active=True).order_by('order')
    return education_levels

def get_grade_financial_summary(grade, current_year):
    """
    حساب ملخص مالي لصف معين باستخدام StudentFinancialService
    """
    students_qs = Student.objects.filter(
        grade_level=grade,
        is_active=True
    ).select_related(
        'grade_level',
        'academic_year'
    )

    students_count = students_qs.count()

    total_fees = Decimal('0.00')
    total_payments = Decimal('0.00')
    total_owed = Decimal('0.00')
    fee_per_student = Decimal('0.00')

    for student in students_qs:
        summary = StudentFinancialService.get_student_balance(
            student,
            current_year
        )

        total_fees += summary['required_fees']
        total_payments += summary['paid_amount']
        total_owed += summary['owed_amount']

        if fee_per_student == 0 and summary['required_fees'] > 0:
            fee_per_student = summary['required_fees']

    return {
        'students_count': students_count,
        'fee_per_student': fee_per_student,
        'total_fees': total_fees,
        'total_payments': total_payments,
        'total_owed': total_owed,
    }


def get_level_financial_summary(level, current_year):
    """
    حساب ملخص مالي لمرحلة تعليمية كاملة باستخدام StudentFinancialService
    """
    students_qs = Student.objects.filter(
        grade_level__education_level=level,
        is_active=True
    ).select_related(
        'grade_level__education_level',
        'academic_year'
    )

    total_students = students_qs.count()

    total_fees = Decimal('0.00')
    total_payments = Decimal('0.00')
    total_owed = Decimal('0.00')

    for student in students_qs:
        summary = StudentFinancialService.get_student_balance(
            student,
            current_year
        )

        total_fees += summary['required_fees']
        total_payments += summary['paid_amount']
        total_owed += summary['owed_amount']

    return {
        'total_students': total_students,
        'total_fees': total_fees,
        'total_payments': total_payments,
        'total_owed': total_owed,
    }
# ===================================
# 🏠 الصفحات الرئيسية
# ===================================

@never_cache
@students_basic_access
def home(request):
    """الصفحة الرئيسية - تحويل للوحة التحكم"""
    return redirect('students:student_affairs_home')


@never_cache
@students_basic_access
def student_affairs_home(request):
    """لوحة تحكم شؤون الطلاب مع الإحصائيات"""
    user_role = get_user_role(request.user)
    system_settings, current_year = get_system_data()

    students_qs = Student.objects.filter(
        is_active=True
    ).select_related(
        'grade_level__education_level',
        'academic_year'
    )

    # إحصائيات عامة
    total_students = students_qs.count()
    male_students = students_qs.filter(gender='M').count()
    female_students = students_qs.filter(gender='F').count()
    students_without_grade = students_qs.filter(grade_level__isnull=True).count()

    # إحصائيات مالية موحدة
    students_paid = 0
    students_owing = 0
    total_outstanding = Decimal('0.00')
    total_required = Decimal('0.00')
    total_paid = Decimal('0.00')
    collection_percentage = 0

    if current_year:
        for student in students_qs:
            balance = StudentFinancialService.get_student_balance(
                student,
                current_year
            )

            total_required += balance['required_fees']
            total_paid += balance['paid_amount']
            total_outstanding += balance['owed_amount']

            if balance['is_paid']:
                students_paid += 1
            else:
                students_owing += 1

        collection_percentage = (
            total_paid / total_required * 100
        ) if total_required > 0 else 0

    # الطلاب المضافون مؤخراً
    recent_students = students_qs.order_by('-created_at')[:10]

    # توزيع الطلاب حسب المراحل التعليمية
    education_levels_stats = []

    for level in EducationLevel.objects.filter(is_active=True).order_by('order', 'name'):
        count = students_qs.filter(
            grade_level__education_level=level
        ).count()

        if count > 0:
            education_levels_stats.append({
                'level': level,
                'student_count': count,
                'percentage': round((count * 100 / total_students) if total_students > 0 else 0, 1),
            })

    # توزيع الطلاب حسب الصفوف
    grade_levels_stats = []

    for grade in GradeLevel.objects.filter(
        is_active=True
    ).select_related(
        'education_level'
    ).order_by(
        'education_level__order',
        'order',
        'name'
    ):
        count = students_qs.filter(grade_level=grade).count()

        if count > 0:
            grade_levels_stats.append({
                'grade': grade,
                'student_count': count,
                'percentage': round((count * 100 / total_students) if total_students > 0 else 0, 1),
            })

    # صلاحيات المستخدم للقالب
    permissions = {
        'can_add': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER', 'STUDENT_AFFAIRS'],
        'can_edit': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER', 'STUDENT_AFFAIRS'],
        'can_delete': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER', 'STUDENT_AFFAIRS'],
        'can_archive': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER', 'STUDENT_AFFAIRS'],

        'can_import': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER', 'STUDENT_AFFAIRS'],
        'can_export': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER', 'STUDENT_AFFAIRS'],

        'can_reports': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER'],
        'can_upgrade': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER'],
        'can_financial_sync': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER'],

        'is_student_affairs_only': user_role == 'STUDENT_AFFAIRS',
        'is_manager': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER'],
        'is_system_admin': user_role == 'SYSTEM_ADMIN',
    }

    context = {
        'system_settings': system_settings,
        'current_year': current_year,

        'total_students': total_students,
        'male_students': male_students,
        'female_students': female_students,
        'students_without_grade': students_without_grade,

        'students_paid': students_paid,
        'students_owing': students_owing,
        'total_outstanding': total_outstanding,
        'total_required': total_required,
        'total_paid': total_paid,
        'collection_percentage': collection_percentage,

        'recent_students': recent_students,
        'education_levels_stats': education_levels_stats,
        'grade_levels_stats': grade_levels_stats,

        'today': timezone.now().date(),
        'permissions': permissions,
        'user_role': user_role,
        'title': 'لوحة تحكم شؤون الطلاب',
    }

    return render(request, 'students/student_affairs_home.html', context)

# ===================================
# 👥 إدارة الطلاب
# ===================================

@never_cache
@students_basic_access
def student_list(request):
    """قائمة الطلاب مع البحث والفلاتر"""
    user_role = get_user_role(request.user)
    system_settings, current_academic_year = get_system_data()

    students = Student.objects.filter(
        is_active=True
    ).select_related(
        'grade_level__education_level',
        'academic_year'
    )

    search_query = request.GET.get('search_query', '').strip()
    education_level_id = request.GET.get('education_level', '').strip()
    grade_level_id = request.GET.get('grade_level', '').strip()
    gender = request.GET.get('gender', '').strip()
    has_balance = request.GET.get('has_balance', '').strip()

    if search_query:
        students = students.filter(
            Q(name__icontains=search_query) |
            Q(national_number__icontains=search_query) |
            Q(phone_number__icontains=search_query) |
            Q(parent_name__icontains=search_query) |
            Q(parent_phone__icontains=search_query)
        )

    if education_level_id:
        students = students.filter(
            grade_level__education_level_id=education_level_id
        )

    if grade_level_id:
        students = students.filter(
            grade_level_id=grade_level_id
        )

    if gender:
        students = students.filter(gender=gender)

    if has_balance == 'paid':
        students = students.filter(total_owed__lte=0)
    elif has_balance == 'owing':
        students = students.filter(total_owed__gt=0)

    students = students.order_by(
        'grade_level__education_level__order',
        'grade_level__order',
        'name'
    )

    paginator = Paginator(students, 25)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    education_levels = EducationLevel.objects.filter(
        is_active=True
    ).order_by('order', 'name')

    grade_levels = GradeLevel.objects.filter(
        is_active=True
    ).select_related(
        'education_level'
    ).order_by(
        'education_level__order',
        'order',
        'name'
    )

    context = {
        'students': page_obj,
        'page_obj': page_obj,
        'total_students': students.count(),
        'education_levels': education_levels,
        'grade_levels': grade_levels,
        'search_query': search_query,
        'selected_education_level': education_level_id,
        'selected_grade_level': grade_level_id,
        'selected_gender': gender,
        'selected_has_balance': has_balance,
        'user_role': user_role,
        'current_academic_year': current_academic_year,
        'title': 'قائمة الطلاب',
    }

    return render(request, 'students/student_list.html', context)

@never_cache
@students_add_only
def add_student(request):
    """إضافة طالب جديد - متاح لموظف شؤون الطلاب"""
    user_role = get_user_role(request.user)
    system_settings, current_academic_year = get_system_data()
    education_levels = get_education_data()

    if request.method == 'POST':
        form = StudentForm(request.POST)

        if form.is_valid():
            try:
                student = form.save(commit=False)

                if current_academic_year and not student.academic_year:
                    student.academic_year = current_academic_year

                student.save()

                messages.success(request, f'تم إضافة الطالب {student.name} بنجاح!')

                if user_role == 'STUDENT_AFFAIRS':
                    return redirect('students:student_affairs_home')

                return redirect('students:student_detail', pk=student.pk)

            except Exception as e:
                messages.error(request, f'حدث خطأ أثناء إضافة الطالب: {str(e)}')
        else:
            messages.error(request, 'يرجى مراجعة البيانات المدخلة وتصحيح الأخطاء')
    else:
        form = StudentForm()

    context = {
        'form': form,
        'current_academic_year': current_academic_year,
        'education_levels': education_levels,
        'title': 'إضافة طالب جديد',
        'user_role': user_role,
    }

    return render(request, 'students/add_student.html', context)

@never_cache
@students_basic_access
def student_detail(request, pk):
    """تفاصيل الطالب"""
    student = get_object_or_404(
        Student.objects.select_related(
            'grade_level__education_level',
            'academic_year'
        ),
        pk=pk
    )

    user_role = get_user_role(request.user)
    system_settings, current_academic_year = get_system_data()

    financial_summary = StudentFinancialService.get_student_balance(
        student,
        current_academic_year
    )

    context = {
        'student': student,
        'user_role': user_role,
        'system_settings': system_settings,
        'current_academic_year': current_academic_year,
        'financial_summary': financial_summary,
        'title': f'تفاصيل الطالب - {student.name}',
    }

    return render(request, 'students/student_detail.html', context)

@never_cache
@students_sensitive_operation
def sync_students_financial_data(request):
    """مزامنة الحقول المالية للطلاب حسب إعدادات المصروفات الحالية"""
    if request.method != 'POST':
        messages.warning(request, 'طريقة الطلب غير صحيحة')
        return redirect('students:student_list')

    system_settings, current_academic_year = get_system_data()

    if not current_academic_year:
        messages.error(request, 'لا يوجد عام دراسي حالي محدد')
        return redirect('students:student_list')

    students = Student.objects.filter(
        is_active=True
    ).select_related(
        'grade_level',
        'academic_year'
    )

    updated_count = 0
    skipped_count = 0

    for student in students:
        if not student.grade_level:
            skipped_count += 1
            continue

        try:
            StudentFinancialService.sync_student_financial_fields(
                student,
                academic_year=current_academic_year,
                save=True
            )
            updated_count += 1
        except Exception:
            skipped_count += 1

    messages.success(
        request,
        f'تمت مزامنة البيانات المالية لعدد {updated_count} طالب. تم تخطي {skipped_count} طالب.'
    )

    return redirect('students:student_list')


@never_cache
@students_edit_access
def edit_student(request, pk):
    """تعديل بيانات الطالب - متاح للمدير وموظف شئون الطلاب"""
    student = get_object_or_404(Student, pk=pk)
    user_role = get_user_role(request.user)
    system_settings, current_academic_year = get_system_data()
    education_levels = get_education_data()

    if request.method == 'POST':
        form = Student_edit_Form(request.POST, instance=student)

        if form.is_valid():
            try:
                student = form.save()
                messages.success(request, 'تم تحديث بيانات الطالب بنجاح!')
                return redirect('students:student_detail', pk=student.pk)

            except Exception as e:
                messages.error(request, f'حدث خطأ أثناء تحديث البيانات: {str(e)}')
        else:
            messages.error(request, 'يرجى مراجعة البيانات المدخلة وتصحيح الأخطاء')
    else:
        form = Student_edit_Form(instance=student)

    context = {
        'form': form,
        'student': student,
        'current_academic_year': current_academic_year,
        'education_levels': education_levels,
        'user_role': user_role,
        'title': f'تعديل بيانات الطالب - {student.name}',
    }

    return render(request, 'students/edit_student_form.html', context)

@never_cache
@students_archive_access
def confirm_delete_student(request, student_id):
    """أرشفة الطالب بدلاً من الحذف النهائي"""
    student = get_object_or_404(
        Student.objects.select_related(
            'grade_level__education_level',
            'academic_year',
        ),
        id=student_id
    )

    user_role = get_user_role(request.user)

    if request.method == 'POST':
        try:
            ArchiveStudent.objects.create(
                # البيانات الأساسية
                archive_name=student.name,
                archive_national_number=student.national_number or '',
                archive_passport_number=getattr(student, 'passport_number', '') or '',
                archive_student_type=student.get_student_type_display() if hasattr(student, 'get_student_type_display') else '',
                archive_nationality=getattr(student, 'nationality', '') or '',
                archive_religion=student.get_religion_display() if hasattr(student, 'get_religion_display') and student.religion else '',

                archive_age=student.age or 0,
                archive_gender=student.gender or '',
                archive_date_of_birth=student.date_of_birth,

                # البيانات الأكاديمية
                archive_academic_year=str(student.academic_year) if student.academic_year else "غير محدد",
                archive_grade_level=student.grade_name if hasattr(student, 'grade_name') else "غير محدد",
                archive_education_level=student.education_level_name if hasattr(student, 'education_level_name') else "غير محدد",
                archive_enrollment_status=student.get_enrollment_status_display() if hasattr(student, 'get_enrollment_status_display') else '',
                archive_transferred_from_school=getattr(student, 'transferred_from_school', '') or '',
                archive_transferred_to_school=getattr(student, 'transferred_to_school', '') or '',

                # بيانات الدمج
                archive_is_integration_student=getattr(student, 'is_integration_student', False),
                archive_disability_type=getattr(student, 'disability_type', '') or '',
                archive_subject_exemptions=student.get_subject_exemptions_display() if hasattr(student, 'get_subject_exemptions_display') else '',

                # البيانات المالية
                archive_total_payments=student.total_payments or Decimal('0.00'),
                archive_total_fees=student.total_fees or Decimal('0.00'),
                archive_total_owed=student.total_owed or Decimal('0.00'),

                # ولي الأمر والولاية التعليمية
                archive_parent_name=student.parent_name or '',
                archive_parent_phone=student.parent_phone or '',
                archive_father_job=getattr(student, 'father_job', '') or '',
                archive_educational_guardian=student.get_educational_guardian_display() if hasattr(student, 'get_educational_guardian_display') else '',
                archive_educational_guardian_name=getattr(student, 'educational_guardian_name', '') or '',

                # أبناء العاملين
                archive_is_staff_child=getattr(student, 'is_staff_child', False),
                archive_staff_parent_name=getattr(student, 'staff_parent_name', '') or '',
                archive_staff_parent_job=getattr(student, 'staff_parent_job', '') or '',

                # بيانات الأرشفة
                archived_reason='أرشفة من شاشة شئون الطلاب',
                archived_by=request.user if request.user.is_authenticated else None,
            )

            student_name = student.name
            student.delete()

            messages.success(
                request,
                f'تم أرشفة الطالب {student_name} بنجاح ونقله إلى أرشيف الطلاب.'
            )
            return redirect('students:student_list')

        except Exception as e:
            messages.error(request, f'حدث خطأ أثناء أرشفة الطالب: {str(e)}')

    context = {
        'student': student,
        'user_role': user_role,
        'title': f'أرشفة الطالب - {student.name}',
        'warning_message': 'سيتم نقل الطالب إلى الأرشيف مع الاحتفاظ ببياناته الأساسية والجديدة.',
    }

    return render(request, 'students/confirm_delete_student.html', context)


# ===================================
# 🔍 البحث والتصفح
# ===================================
@never_cache
@students_basic_access
def search_student(request):
    """صفحة البحث المتقدم"""
    user_role = get_user_role(request.user)

    education_levels = EducationLevel.objects.filter(
        is_active=True
    ).order_by(
        'order',
        'name'
    )

    grade_levels = GradeLevel.objects.filter(
        is_active=True
    ).select_related(
        'education_level'
    ).order_by(
        'education_level__order',
        'order',
        'name'
    )

    context = {
        'user_role': user_role,
        'education_levels': education_levels,
        'grade_levels': grade_levels,
        'title': 'البحث المتقدم عن الطلاب',
        
        'can_add_student': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER', 'STUDENT_AFFAIRS'],
        'can_edit_student': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER', 'STUDENT_AFFAIRS'],
        'can_delete_student': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER', 'STUDENT_AFFAIRS'],
        'can_archive_student': user_role in ['SYSTEM_ADMIN', 'SCHOOL_MANAGER', 'STUDENT_AFFAIRS'],

    }

    return render(request, 'students/search_student.html', context)


@never_cache
@students_basic_access
def ajax_student_search(request):
    """بحث AJAX عن الطلاب مع دعم الفلاتر"""
    query = request.GET.get('q', '').strip()
    education_level = request.GET.get('education_level', '').strip()
    grade_level = request.GET.get('grade_level', '').strip()
    gender = request.GET.get('gender', '').strip()
    age_from = request.GET.get('age_from', '').strip()
    age_to = request.GET.get('age_to', '').strip()

    if (
        len(query) < 2
        and not education_level
        and not grade_level
        and not gender
        and not age_from
        and not age_to
    ):
        return JsonResponse({
            'success': True,
            'students': [],
            'count': 0,
        })

    students = Student.objects.filter(
        is_active=True
    ).select_related(
        'grade_level__education_level',
        'academic_year'
    )

    if query:
        students = students.filter(
            Q(name__icontains=query) |
            Q(national_number__icontains=query) |
            Q(phone_number__icontains=query) |
            Q(parent_name__icontains=query) |
            Q(parent_phone__icontains=query)
        )

    if education_level:
        students = students.filter(
            grade_level__education_level_id=education_level
        )

    if grade_level:
        students = students.filter(
            grade_level_id=grade_level
        )

    if gender:
        students = students.filter(
            gender=gender
        )

    if age_from:
        try:
            students = students.filter(age__gte=int(age_from))
        except ValueError:
            pass

    if age_to:
        try:
            students = students.filter(age__lte=int(age_to))
        except ValueError:
            pass

    students = students.order_by(
        'grade_level__education_level__order',
        'grade_level__order',
        'name'
    )[:50]

    students_data = []

    for student in students:
        grade_name = ''
        education_level_name = ''

        if student.grade_level:
            grade_name = student.grade_level.name

            if student.grade_level.education_level:
                education_level_name = student.grade_level.education_level.name

        students_data.append({
            'id': student.id,
            'name': student.name or '',
            'national_number': student.national_number or '',
            'phone_number': student.phone_number or '',
            'parent_name': student.parent_name or '',
            'parent_phone': student.parent_phone or '',
            'age': student.age or '',
            'gender': student.gender or '',
            'gender_display': 'ذكر' if student.gender == 'M' else 'أنثى' if student.gender == 'F' else 'غير محدد',
            'gender_icon': 'male' if student.gender == 'M' else 'female' if student.gender == 'F' else 'question',
            'grade_level': grade_name,
            'education_level': education_level_name,
            'academic_year': student.academic_year.name if student.academic_year else '',
            'total_fees': float(student.total_fees or 0),
            'total_payments': float(student.total_payments or 0),
            'total_owed': float(student.total_owed or 0),
            'financial_status': student.get_financial_status() if hasattr(student, 'get_financial_status') else '',
            'status_color': 'danger' if (student.total_owed or 0) > 0 else 'success',
            'detail_url': reverse('students:student_detail', kwargs={'pk': student.pk}),
            'edit_url': reverse('students:edit_student', kwargs={'pk': student.pk}),
            'delete_url': reverse('students:confirm_delete_student', kwargs={'student_id': student.pk}),
        })

    return JsonResponse({
        'success': True,
        'students': students_data,
        'count': len(students_data),
        'query': query,
    })

# ===================================
# 🔌 APIs المساعدة
# ===================================

@never_cache
@students_basic_access
def get_grades_by_level(request, level_id):
    """API لجلب الصفوف الدراسية حسب المرحلة التعليمية"""
    try:
        education_level = get_object_or_404(EducationLevel, id=level_id, is_active=True)
        grades = GradeLevel.objects.filter(
            education_level=education_level, 
            is_active=True
        ).order_by('order')
        
        grades_data = []
        for grade in grades:
            grades_data.append({
                'id': grade.id,
                'name': grade.name,
                'name_en': grade.name_en or '',
                'typical_age': grade.typical_age,
                'grade_number': grade.grade_number,
            })
        
        return JsonResponse({
            'success': True,
            'grades': grades_data
        })
        
    except Exception as e:
        return JsonResponse({
            'success': False,
            'error': str(e)
        })

# ===================================
# 📊 التقارير (حسب الصلاحية)
# ===================================

@never_cache
@students_reports_access
def report(request):
    """التقرير العام للطلاب"""
    user_role = get_user_role(request.user)
    system_settings, current_year = get_system_data()

    students_qs = Student.objects.filter(
        is_active=True
    ).select_related(
        'grade_level__education_level',
        'academic_year'
    )

    total_students = students_qs.count()
    male_students = students_qs.filter(gender='M').count()
    female_students = students_qs.filter(gender='F').count()
    students_without_grade = students_qs.filter(grade_level__isnull=True).count()

    # =========================
    # إحصائيات المراحل التعليمية
    # =========================
    education_level_stats = []

    for level in EducationLevel.objects.filter(is_active=True).order_by('order', 'name'):
        count = students_qs.filter(
            grade_level__education_level=level
        ).count()

        if count > 0:
            education_level_stats.append({
                'level': level,
                'count': count,
                'percentage': round((count * 100 / total_students) if total_students > 0 else 0, 1),
            })

    # =========================
    # إحصائيات الصفوف الدراسية
    # =========================
    grade_level_stats = []

    for grade in GradeLevel.objects.filter(
        is_active=True
    ).select_related(
        'education_level'
    ).order_by(
        'education_level__order',
        'order',
        'name'
    ):
        count = students_qs.filter(
            grade_level=grade
        ).count()

        if count > 0:
            grade_level_stats.append({
                'grade': grade,
                'count': count,
                'percentage': round((count * 100 / total_students) if total_students > 0 else 0, 1),
            })

    # =========================
    # إحصائيات مالية
    # =========================
    financial_summary = {}

    if user_role == 'SYSTEM_ADMIN' and current_year:
        total_required = Decimal('0.00')
        total_paid = Decimal('0.00')
        total_owed = Decimal('0.00')
        paid_students = 0
        owing_students = 0

        for student in students_qs:
            balance = StudentFinancialService.get_student_balance(
                student,
                current_year
            )

            total_required += balance['required_fees']
            total_paid += balance['paid_amount']
            total_owed += balance['owed_amount']

            if balance['is_paid']:
                paid_students += 1
            else:
                owing_students += 1

        financial_summary = {
            'total_required': total_required,
            'total_paid': total_paid,
            'total_owed': total_owed,
            'paid_students': paid_students,
            'owing_students': owing_students,
            'collection_percentage': (
                total_paid / total_required * 100
            ) if total_required > 0 else 0,
        }

    context = {
        'system_settings': system_settings,
        'current_year': current_year,
        'user_role': user_role,
        'title': 'التقرير العام للطلاب',

        'total_students': total_students,
        'male_students': male_students,
        'female_students': female_students,
        'students_without_grade': students_without_grade,

        'education_level_stats': education_level_stats,
        'grade_level_stats': grade_level_stats,
        'financial_summary': financial_summary,
    }

    return render(request, 'students/report.html', context)

@never_cache
@students_reports_access
def all_reports(request):
    """التقارير الشاملة للطلاب"""
    user_role = get_user_role(request.user)
    system_settings, current_year = get_system_data()

    today = timezone.now().date()
    week_start = today - timedelta(days=7)
    month_start = today.replace(day=1)
    year_start = today.replace(month=1, day=1)

    active_students = Student.objects.filter(
        is_active=True
    ).select_related(
        'grade_level__education_level',
        'academic_year'
    )

    total_students = active_students.count()
    male_students = active_students.filter(gender='M').count()
    female_students = active_students.filter(gender='F').count()

    time_stats = {
        'today': active_students.filter(created_at__date=today).count(),
        'this_week': active_students.filter(created_at__date__gte=week_start).count(),
        'this_month': active_students.filter(created_at__date__gte=month_start).count(),
        'this_year': active_students.filter(created_at__date__gte=year_start).count(),
    }

    # =========================
    # الإحصائيات المالية العامة
    # =========================
    financial_stats = {}

    if user_role == 'SYSTEM_ADMIN' and current_year:
        try:
            total_calculated_fees = Decimal('0.00')
            total_payments = Decimal('0.00')
            total_owed = Decimal('0.00')
            paid_students = 0
            unpaid_students = 0

            for student in active_students:
                summary = StudentFinancialService.get_student_balance(
                    student,
                    current_year
                )

                total_calculated_fees += summary['required_fees']
                total_payments += summary['paid_amount']
                total_owed += summary['owed_amount']

                if summary['is_paid']:
                    paid_students += 1
                else:
                    unpaid_students += 1

            total_fees_from_settings = SchoolFeesSettings.objects.filter(
                academic_year=current_year,
                is_active=True
            ).aggregate(
                total_fees=Sum('total_amount')
            )['total_fees'] or Decimal('0.00')

            financial_stats = {
                'total_fees_from_settings': total_fees_from_settings,
                'total_calculated_fees': total_calculated_fees,
                'total_payments': total_payments,
                'total_owed': total_owed,
                'paid_students': paid_students,
                'unpaid_students': unpaid_students,
                'collection_percentage': (
                    total_payments / total_calculated_fees * 100
                ) if total_calculated_fees > 0 else 0,
            }

        except Exception as e:
            print(f"خطأ في حساب الإحصائيات المالية: {e}")
            financial_stats = {
                'total_fees_from_settings': Decimal('0.00'),
                'total_calculated_fees': Decimal('0.00'),
                'total_payments': Decimal('0.00'),
                'total_owed': Decimal('0.00'),
                'paid_students': 0,
                'unpaid_students': 0,
                'collection_percentage': 0,
            }

    # =========================
    # إحصائيات إعدادات المصروفات
    # =========================
    fees_settings_stats = {}

    if user_role == 'SYSTEM_ADMIN' and current_year:
        fees_qs = SchoolFeesSettings.objects.filter(
            academic_year=current_year,
            is_active=True
        )

        fees_settings_stats = fees_qs.aggregate(
            total_fee_types=Count('id'),
            average_fee_per_grade=Avg('total_amount'),
            max_fee=Max('total_amount'),
            min_fee=Min('total_amount'),
        )

        fees_settings_stats['total_fee_types'] = fees_settings_stats['total_fee_types'] or 0
        fees_settings_stats['average_fee_per_grade'] = fees_settings_stats['average_fee_per_grade'] or Decimal('0.00')
        fees_settings_stats['max_fee'] = fees_settings_stats['max_fee'] or Decimal('0.00')
        fees_settings_stats['min_fee'] = fees_settings_stats['min_fee'] or Decimal('0.00')
        fees_settings_stats['grades_with_fees'] = fees_qs.values('grade_level').distinct().count()

    # =========================
    # إحصائيات المراحل التعليمية
    # =========================
    education_levels_stats = []

    for level in EducationLevel.objects.filter(is_active=True).order_by('order', 'name'):
        level_students_qs = active_students.filter(
            grade_level__education_level=level
        )

        students_count = level_students_qs.count()
        level_male = level_students_qs.filter(gender='M').count()
        level_female = level_students_qs.filter(gender='F').count()

        level_financial = {}

        if user_role == 'SYSTEM_ADMIN' and current_year:
            try:
                level_financial = get_level_financial_summary(
                    level,
                    current_year
                )
            except Exception as e:
                print(f"خطأ في حساب مالية المرحلة {level.name}: {e}")
                level_financial = {
                    'total_students': students_count,
                    'total_fees': Decimal('0.00'),
                    'total_payments': Decimal('0.00'),
                    'total_owed': Decimal('0.00'),
                }

        education_levels_stats.append({
            'level': level,
            'total_students': students_count,
            'male_students': level_male,
            'female_students': level_female,
            'percentage': round((students_count * 100 / total_students) if total_students > 0 else 0, 1),
            'financial': level_financial,
        })

    # =========================
    # إحصائيات الصفوف الدراسية
    # =========================
    grade_levels_stats = []

    for grade in GradeLevel.objects.filter(
        is_active=True
    ).select_related(
        'education_level'
    ).order_by(
        'education_level__order',
        'order',
        'name'
    ):
        grade_students_qs = active_students.filter(
            grade_level=grade
        )

        students_count = grade_students_qs.count()

        if students_count == 0:
            continue

        grade_financial = {}

        if user_role == 'SYSTEM_ADMIN' and current_year:
            try:
                grade_financial = get_grade_financial_summary(
                    grade,
                    current_year
                )
            except Exception as e:
                print(f"خطأ في حساب مالية الصف {grade.name}: {e}")
                grade_financial = {
                    'students_count': students_count,
                    'fee_per_student': Decimal('0.00'),
                    'total_fees': Decimal('0.00'),
                    'total_payments': Decimal('0.00'),
                    'total_owed': Decimal('0.00'),
                }

        grade_levels_stats.append({
            'grade': grade,
            'total_students': students_count,
            'education_level': grade.education_level.name if grade.education_level else 'غير محدد',
            'financial': grade_financial,
        })

    context = {
        'system_settings': system_settings,
        'current_year': current_year,
        'report_date': today,
        'user_role': user_role,

        'total_students': total_students,
        'male_students': male_students,
        'female_students': female_students,

        'time_stats': time_stats,
        'financial_stats': financial_stats,
        'fees_settings_stats': fees_settings_stats,
        'education_levels_stats': education_levels_stats,
        'grade_levels_stats': grade_levels_stats,

        'title': 'التقارير الشاملة',
    }

    return render(request, 'students/all_reports.html', context)


@never_cache
@students_reports_access
def daily_report(request):
    """التقرير اليومي للطلاب"""
    user_role = get_user_role(request.user)
    system_settings, current_year = get_system_data()

    selected_date_str = request.GET.get('date', '').strip()

    if selected_date_str:
        try:
            selected_date = datetime.strptime(selected_date_str, '%Y-%m-%d').date()
        except ValueError:
            selected_date = timezone.now().date()
            messages.warning(request, 'صيغة التاريخ غير صحيحة، تم عرض تقرير اليوم الحالي')
    else:
        selected_date = timezone.now().date()

    day_start = selected_date
    day_end = selected_date

    students_qs = Student.objects.filter(
        is_active=True
    ).select_related(
        'grade_level__education_level',
        'academic_year'
    )

    new_students_today = students_qs.filter(
        created_at__date=selected_date
    ).order_by(
        '-created_at'
    )

    updated_students_today = students_qs.filter(
        updated_at__date=selected_date
    ).exclude(
        created_at__date=selected_date
    ).order_by(
        '-updated_at'
    )

    total_students = students_qs.count()
    total_new_today = new_students_today.count()
    total_updated_today = updated_students_today.count()

    male_new_today = new_students_today.filter(gender='M').count()
    female_new_today = new_students_today.filter(gender='F').count()

    # =========================
    # توزيع الطلاب الجدد حسب المرحلة
    # =========================
    education_level_stats = []

    for level in EducationLevel.objects.filter(is_active=True).order_by('order', 'name'):
        count = new_students_today.filter(
            grade_level__education_level=level
        ).count()

        if count > 0:
            education_level_stats.append({
                'level': level,
                'count': count,
                'percentage': round((count * 100 / total_new_today) if total_new_today > 0 else 0, 1),
            })

    # =========================
    # توزيع الطلاب الجدد حسب الصف
    # =========================
    grade_level_stats = []

    for grade in GradeLevel.objects.filter(
        is_active=True
    ).select_related(
        'education_level'
    ).order_by(
        'education_level__order',
        'order',
        'name'
    ):
        count = new_students_today.filter(
            grade_level=grade
        ).count()

        if count > 0:
            grade_level_stats.append({
                'grade': grade,
                'count': count,
                'percentage': round((count * 100 / total_new_today) if total_new_today > 0 else 0, 1),
            })

    # =========================
    # الملخص المالي لطلاب اليوم
    # =========================
    financial_summary = {}

    if user_role == 'SYSTEM_ADMIN' and current_year:
        total_required = Decimal('0.00')
        total_paid = Decimal('0.00')
        total_owed = Decimal('0.00')

        for student in new_students_today:
            balance = StudentFinancialService.get_student_balance(
                student,
                current_year
            )

            total_required += balance['required_fees']
            total_paid += balance['paid_amount']
            total_owed += balance['owed_amount']

        financial_summary = {
            'total_required': total_required,
            'total_paid': total_paid,
            'total_owed': total_owed,
            'collection_percentage': (
                total_paid / total_required * 100
            ) if total_required > 0 else 0,
        }

    context = {
        'system_settings': system_settings,
        'current_year': current_year,
        'user_role': user_role,
        'title': 'التقرير اليومي',

        'selected_date': selected_date,
        'selected_date_str': selected_date.strftime('%Y-%m-%d'),

        'total_students': total_students,
        'total_new_today': total_new_today,
        'total_updated_today': total_updated_today,
        'male_new_today': male_new_today,
        'female_new_today': female_new_today,

        'new_students_today': new_students_today,
        'updated_students_today': updated_students_today,

        'education_level_stats': education_level_stats,
        'grade_level_stats': grade_level_stats,
        'financial_summary': financial_summary,
    }

    return render(request, 'students/daily_report.html', context)


def get_arabic_day_name(english_day):
    """تحويل اسم اليوم من الإنجليزية للعربية"""
    days_map = {
        'Monday': 'الإثنين',
        'Tuesday': 'الثلاثاء', 
        'Wednesday': 'الأربعاء',
        'Thursday': 'الخميس',
        'Friday': 'الجمعة',
        'Saturday': 'السبت',
        'Sunday': 'الأحد'
    }
    return days_map.get(english_day, english_day)

                  
@never_cache
@students_reports_access
def student_dashboard(request):
    """لوحة إحصائيات الطلاب"""
    user_role = get_user_role(request.user)
    system_settings, current_year = get_system_data()

    today = timezone.now().date()
    month_start = today.replace(day=1)
    year_start = today.replace(month=1, day=1)

    students_qs = Student.objects.filter(
        is_active=True
    ).select_related(
        'grade_level__education_level',
        'academic_year'
    )

    total_students = students_qs.count()
    male_students = students_qs.filter(gender='M').count()
    female_students = students_qs.filter(gender='F').count()

    new_this_month = students_qs.filter(
        created_at__date__gte=month_start
    ).count()

    new_this_year = students_qs.filter(
        created_at__date__gte=year_start
    ).count()

    students_without_grade = students_qs.filter(
        grade_level__isnull=True
    ).count()

    # =========================
    # توزيع الطلاب على المراحل
    # =========================
    education_distribution = []

    for level in EducationLevel.objects.filter(is_active=True).order_by('order', 'name'):
        count = students_qs.filter(
            grade_level__education_level=level
        ).count()

        if count > 0:
            education_distribution.append({
                'level': level,
                'count': count,
                'percentage': round((count * 100 / total_students) if total_students > 0 else 0, 1),
            })

    # =========================
    # توزيع الطلاب على الصفوف
    # =========================
    grade_distribution = []

    for grade in GradeLevel.objects.filter(
        is_active=True
    ).select_related(
        'education_level'
    ).order_by(
        'education_level__order',
        'order',
        'name'
    ):
        count = students_qs.filter(grade_level=grade).count()

        if count > 0:
            grade_distribution.append({
                'grade': grade,
                'count': count,
                'percentage': round((count * 100 / total_students) if total_students > 0 else 0, 1),
            })

    # =========================
    # إحصائيات مالية للمدير العام فقط
    # =========================
    financial_summary = {}

    if user_role == 'SYSTEM_ADMIN' and current_year:
        total_required = Decimal('0.00')
        total_paid = Decimal('0.00')
        total_owed = Decimal('0.00')
        paid_students = 0
        owing_students = 0

        for student in students_qs:
            balance = StudentFinancialService.get_student_balance(
                student,
                current_year
            )

            total_required += balance['required_fees']
            total_paid += balance['paid_amount']
            total_owed += balance['owed_amount']

            if balance['is_paid']:
                paid_students += 1
            else:
                owing_students += 1

        financial_summary = {
            'total_required': total_required,
            'total_paid': total_paid,
            'total_owed': total_owed,
            'paid_students': paid_students,
            'owing_students': owing_students,
            'collection_percentage': (
                total_paid / total_required * 100
            ) if total_required > 0 else 0,
        }

    context = {
        'system_settings': system_settings,
        'current_year': current_year,
        'user_role': user_role,
        'title': 'لوحة إحصائيات الطلاب',

        'total_students': total_students,
        'male_students': male_students,
        'female_students': female_students,
        'new_this_month': new_this_month,
        'new_this_year': new_this_year,
        'students_without_grade': students_without_grade,

        'education_distribution': education_distribution,
        'grade_distribution': grade_distribution,
        'financial_summary': financial_summary,
    }

    return render(request, 'students/student_dashboard.html', context)

# ===================================
# 🔧 الأدوات الإدارية (مدير عام فقط)
# ===================================

@never_cache
@students_import_export_access
def export_students(request):
    """تصدير بيانات الطلاب - للمدير العام فقط"""
    export_format = (
        request.POST.get('export_format')
        or request.POST.get('format')
        or request.GET.get('export_format')
        or request.GET.get('format')
        or 'csv'
    )

    include_inactive = (
        request.POST.get('include_inactive') == 'on'
        or request.GET.get('include_inactive') == '1'
    )

    selected_grade_levels = request.POST.getlist('grade_levels') or request.GET.getlist('grade_levels')

    queryset = Student.objects.all().select_related(
        'grade_level__education_level',
        'academic_year'
    )

    if not include_inactive:
        queryset = queryset.filter(is_active=True)

    if selected_grade_levels:
        queryset = queryset.filter(grade_level_id__in=selected_grade_levels)

    queryset = queryset.order_by(
        'grade_level__education_level__order',
        'grade_level__order',
        'name'
    )

    if request.method == 'POST' or request.GET.get('download') == '1':
        try:
            return StudentExportService.export(queryset, export_format)
        except Exception as e:
            messages.error(request, f'حدث خطأ في التصدير: {str(e)}')

    grade_levels = GradeLevel.objects.filter(
        is_active=True
    ).select_related(
        'education_level'
    ).order_by(
        'education_level__order',
        'order',
        'name'
    )

    context = {
        'total_students': Student.objects.filter(is_active=True).count(),
        'grade_levels': grade_levels,
        'title': 'تصدير بيانات الطلاب',
        'export_formats': [
            {'value': 'csv', 'label': 'CSV'},
            {'value': 'excel', 'label': 'Excel'},
            {'value': 'json', 'label': 'JSON'},
        ],
    }

    return render(request, 'students/export_students.html', context)

@never_cache
@students_sensitive_operation
def upgrade_students(request):
    """تحويل المسار القديم إلى معالج الترقية الجديد"""
    return redirect('students:upgrade_students_wizard')


@never_cache
@students_import_export_access
def export_students_advanced(request):
    """
    تحويل التصدير المتقدم إلى صفحة التصدير الموحدة
    لتجنب تكرار الكود والشاشات
    """
    return redirect('students:export_students')


# استبدل دالة download_import_template في students/views.py بهذه النسخة

# @never_cache
# @students_import_export_access
# def download_import_template(request):
#     """تحميل قالب استيراد الطلاب CSV أو Excel متوافق مع الحقول الجديدة"""
#     template_format = request.GET.get('format', 'excel').lower()

#     headers = [
#     'الاسم',
#     'نوع الطالب',
#     'الرقم القومي',
#     'رقم جواز السفر',
#     'الجنسية',
#     'الديانة',
#     'العمر',
#     'النوع',
#     'تاريخ الميلاد',
#     'رقم الهاتف',
#     'العنوان',

#     'العام الدراسي',
#     'المرحلة التعليمية',
#     'الصف الدراسي',
#     'حالة القيد',
#     'محول من مدرسة',
#     'محول إلى مدرسة',

#     'طالب دمج',
#     'نوع الإعاقة',
#     'إعفاء من العربي',
#     'إعفاء من الإنجليزي',
#     'إعفاء من الفرنسي',
#     'إعفاءات أخرى',

#     'من أبناء العاملين',
#     'اسم الموظف',
#     'وظيفة الموظف داخل المدرسة',

#     'اسم ولي الأمر',
#     'هاتف ولي الأمر',
#     'بريد ولي الأمر',
#     'وظيفة الأب',
#     'صاحب الولاية التعليمية',
#     'اسم صاحب الولاية التعليمية',
#     'هاتف صاحب الولاية التعليمية',
#     ]

#     sample_rows = [
#         [
#             'أحمد محمد علي',
#             'طالب عادي',
#             '30001011234567',
#             '',
#             'مصري',
#             'مسلم',
#             '',
#             'ذكر',
#             '',
#             '01000000000',
#             'القاهرة',

#             '',
#             'الابتدائية',
#             'الأول الابتدائي',
#             'مستجد',
#             '',
#             '',

#             'لا',
#             '',
#             'لا',
#             'لا',
#             'لا',
#             '',

#             'لا',
#             '',
#             '',

#             'محمد علي',
#             '01111111111',
#             'parent@example.com',
#             'محاسب',
#             'الأب',
#             '',
#             '',
#         ],
#         [
#             'يوسف جون سميث',
#             'وافد',
#             '',
#             'A12345678',
#             'أمريكي',
#             'مسيحي',
#             '8',
#             'ذكر',
#             '2018-05-10',
#             '01000000002',
#             'الإسماعيلية',

#             '',
#             'الثاني الابتدائي',
#             'محول',
#             'مدرسة دولية سابقة',
#             '',

#             'لا',
#             '',
#             'لا',
#             'لا',
#             'لا',
#             '',

#             'لا',
#             '',
#             '',

#             'John Smith',
#             '01111111113',
#             'john@example.com',
#             'مهندس',
#             'الأب',
#             '',
#             '',
#         ],
#         [
#             'مريم أحمد حسن',
#             'طالب عادي',
#             '30303031234567',
#             '',
#             'مصري',
#             'مسلم',
#             '',
#             'أنثى',
#             '',
#             '01000000003',
#             'القاهرة',

#             '',
#             'الثالث الابتدائي',
#             'ناجح ومنقول',
#             '',
#             '',

#             'نعم',
#             'صعوبات تعلم',
#             'لا',
#             'لا',
#             'نعم',
#             'إعفاء من الفرنسي',

#             'نعم',
#             'أحمد حسن',
#             'مدرس لغة عربية',

#             'أحمد حسن',
#             '01111111114',
#             '',
#             'مدرس',
#             'الأم',
#             '',
#             '',
#         ],
#     ]

#     if template_format == 'csv':
#         response = HttpResponse(content_type='text/csv; charset=utf-8-sig')
#         response['Content-Disposition'] = 'attachment; filename="student_import_template_extended.csv"'
#         response.write('\ufeff')

#         writer = csv.writer(response)
#         writer.writerow(headers)
#         writer.writerows(sample_rows)

#         return response

#     workbook = openpyxl.Workbook()
#     worksheet = workbook.active
#     worksheet.title = 'قالب استيراد الطلاب'
#     worksheet.sheet_view.rightToLeft = True

#     header_fill = PatternFill(
#         start_color='1F4E79',
#         end_color='1F4E79',
#         fill_type='solid'
#     )
#     header_font = Font(
#         bold=True,
#         color='FFFFFF'
#     )
#     center_alignment = Alignment(
#         horizontal='center',
#         vertical='center',
#         wrap_text=True
#     )

#     for col_num, header in enumerate(headers, 1):
#         cell = worksheet.cell(row=1, column=col_num, value=header)
#         cell.fill = header_fill
#         cell.font = header_font
#         cell.alignment = center_alignment

#     for row_num, row_data in enumerate(sample_rows, 2):
#         for col_num, value in enumerate(row_data, 1):
#             cell = worksheet.cell(row=row_num, column=col_num, value=value)
#             cell.alignment = center_alignment

#     instructions = [
#         'تعليمات الاستيراد:',
#         '1. الاسم هو الحقل الإجباري الوحيد.',
#         '2. الرقم القومي اختياري، وإذا تم إدخاله يجب أن يكون 14 رقم وغير مكرر.',
#         '3. رقم جواز السفر اختياري ويستخدم غالباً للوافدين، ويجب ألا يكون مكرراً إذا تم إدخاله.',
#         '4. نوع الطالب يقبل: طالب عادي / وافد.',
#         '5. النوع يقبل: ذكر / أنثى / M / F.',
#         '6. حالة القيد تقبل: مستجد / ناجح ومنقول / محول / باق للإعادة.',
#         '7. صاحب الولاية التعليمية يقبل: الأب / الأم / آخر.',
#         '8. حقول نعم/لا تقبل: نعم / لا / 1 / 0.',
#         '9. الصف الدراسي يمكن كتابته بالاسم مثل: الأول الابتدائي.',
#         '10. العام الدراسي يمكن تركه فارغاً وسيتم استخدام العام الحالي إن وجد.',
#         '11. تاريخ الميلاد اختياري ويفضل بصيغة YYYY-MM-DD.',
#         '12. لا تغير أسماء الأعمدة في الصف الأول.',
#     ]

#     start_row = len(sample_rows) + 4

#     for index, instruction in enumerate(instructions):
#         cell = worksheet.cell(row=start_row + index, column=1, value=instruction)

#         if index == 0:
#             cell.font = Font(bold=True, color='D32F2F')
#         else:
#             cell.font = Font(color='666666')

#         cell.alignment = Alignment(horizontal='right', vertical='center')

#     worksheet.freeze_panes = 'A2'
#     worksheet.auto_filter.ref = f'A1:{worksheet.cell(row=1, column=len(headers)).coordinate}'

#     for column_cells in worksheet.columns:
#         max_length = 0
#         column_letter = column_cells[0].column_letter

#         for cell in column_cells:
#             try:
#                 cell_length = len(str(cell.value)) if cell.value else 0
#                 if cell_length > max_length:
#                     max_length = cell_length
#             except Exception:
#                 pass

#         worksheet.column_dimensions[column_letter].width = min(max_length + 4, 35)

#     output = io.BytesIO()
#     workbook.save(output)
#     output.seek(0)

#     response = HttpResponse(
#         output.read(),
#         content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
#     )
#     response['Content-Disposition'] = 'attachment; filename="student_import_template_extended.xlsx"'

#     return response

@never_cache
@students_import_export_access
def download_import_template(request):
    """تحميل قالب استيراد الطلاب CSV أو Excel متوافق مع الحقول الجديدة"""
    template_format = request.GET.get('format', 'excel').lower()

    headers = [
        'الاسم',
        'نوع الطالب',
        'الرقم القومي',
        'رقم جواز السفر',
        'الجنسية',
        'الديانة',
        'العمر',
        'النوع',
        'تاريخ الميلاد',
        'رقم الهاتف',
        'العنوان',

        'العام الدراسي',
        'المرحلة التعليمية',
        'الصف الدراسي',
        'حالة القيد',
        'محول من مدرسة',
        'محول إلى مدرسة',

        'طالب دمج',
        'نوع الإعاقة',
        'إعفاء من العربي',
        'إعفاء من الإنجليزي',
        'إعفاء من الفرنسي',
        'إعفاءات أخرى',

        'من أبناء العاملين',
        'اسم الموظف',
        'وظيفة الموظف داخل المدرسة',

        'اسم ولي الأمر',
        'هاتف ولي الأمر',
        'بريد ولي الأمر',
        'وظيفة الأب',
        'صاحب الولاية التعليمية',
        'اسم صاحب الولاية التعليمية',
        'هاتف صاحب الولاية التعليمية',
    ]

    sample_rows = [
        [
            'أحمد محمد علي',
            'طالب عادي',
            '30001011234567',
            '',
            'مصري',
            'مسلم',
            '',
            'ذكر',
            '',
            '01000000000',
            'القاهرة',

            '',
            'الابتدائية',
            'الأول الابتدائي',
            'مستجد',
            '',
            '',

            'لا',
            '',
            'لا',
            'لا',
            'لا',
            '',

            'لا',
            '',
            '',

            'محمد علي',
            '01111111111',
            'parent@example.com',
            'محاسب',
            'الأب',
            '',
            '',
        ],
        [
            'يوسف جون سميث',
            'وافد',
            '',
            'A12345678',
            'أمريكي',
            'مسيحي',
            '8',
            'ذكر',
            '2018-05-10',
            '01000000002',
            'الإسماعيلية',

            '',
            'الابتدائية',
            'الثاني الابتدائي',
            'محول',
            'مدرسة دولية سابقة',
            '',

            'لا',
            '',
            'لا',
            'لا',
            'لا',
            '',

            'لا',
            '',
            '',

            'John Smith',
            '01111111113',
            'john@example.com',
            'مهندس',
            'الأب',
            '',
            '',
        ],
        [
            'مريم أحمد حسن',
            'طالب عادي',
            '30303031234567',
            '',
            'مصري',
            'مسلم',
            '',
            'أنثى',
            '',
            '01000000003',
            'القاهرة',

            '',
            'الابتدائية',
            'الثالث الابتدائي',
            'ناجح ومنقول',
            '',
            '',

            'نعم',
            'صعوبات تعلم',
            'لا',
            'لا',
            'نعم',
            'إعفاء من الفرنسي',

            'نعم',
            'أحمد حسن',
            'مدرس لغة عربية',

            'أحمد حسن',
            '01111111114',
            '',
            'مدرس',
            'الأم',
            '',
            '',
        ],
    ]

    if template_format == 'csv':
        response = HttpResponse(content_type='text/csv; charset=utf-8-sig')
        response['Content-Disposition'] = 'attachment; filename="student_import_template_extended.csv"'
        response.write('\ufeff')

        writer = csv.writer(response)
        writer.writerow(headers)
        writer.writerows(sample_rows)

        return response

    workbook = openpyxl.Workbook()
    worksheet = workbook.active
    worksheet.title = 'قالب استيراد الطلاب'
    worksheet.sheet_view.rightToLeft = True

    header_fill = PatternFill(
        start_color='1F4E79',
        end_color='1F4E79',
        fill_type='solid'
    )
    header_font = Font(
        bold=True,
        color='FFFFFF'
    )
    center_alignment = Alignment(
        horizontal='center',
        vertical='center',
        wrap_text=True
    )

    for col_num, header in enumerate(headers, 1):
        cell = worksheet.cell(row=1, column=col_num, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_alignment

    for row_num, row_data in enumerate(sample_rows, 2):
        for col_num, value in enumerate(row_data, 1):
            cell = worksheet.cell(row=row_num, column=col_num, value=value)
            cell.alignment = center_alignment

    instructions = [
        'تعليمات الاستيراد:',
        '1. الاسم هو الحقل الإجباري الوحيد.',
        '2. الرقم القومي اختياري، وإذا تم إدخاله يجب أن يكون 14 رقم وغير مكرر.',
        '3. رقم جواز السفر اختياري ويستخدم غالباً للوافدين، ويجب ألا يكون مكرراً إذا تم إدخاله.',
        '4. نوع الطالب يقبل: طالب عادي / وافد.',
        '5. النوع يقبل: ذكر / أنثى / M / F.',
        '6. حالة القيد تقبل: مستجد / ناجح ومنقول / محول / باق للإعادة.',
        '7. صاحب الولاية التعليمية يقبل: الأب / الأم / آخر.',
        '8. حقول نعم/لا تقبل: نعم / لا / 1 / 0.',
        '9. المرحلة التعليمية يمكن كتابتها بالاسم مثل: الابتدائية / الإعدادية / الثانوية.',
        '10. الصف الدراسي يمكن كتابته بالاسم مثل: الأول الابتدائي.',
        '11. عند تكرار اسم الصف في أكثر من مرحلة، يجب كتابة المرحلة التعليمية.',
        '12. العام الدراسي يمكن تركه فارغاً وسيتم استخدام العام الحالي إن وجد.',
        '13. تاريخ الميلاد اختياري ويفضل بصيغة YYYY-MM-DD.',
        '14. لا تغير أسماء الأعمدة في الصف الأول.',
    ]

    start_row = len(sample_rows) + 4

    for index, instruction in enumerate(instructions):
        cell = worksheet.cell(row=start_row + index, column=1, value=instruction)

        if index == 0:
            cell.font = Font(bold=True, color='D32F2F')
        else:
            cell.font = Font(color='666666')

        cell.alignment = Alignment(horizontal='right', vertical='center')

    worksheet.freeze_panes = 'A2'
    worksheet.auto_filter.ref = f'A1:{worksheet.cell(row=1, column=len(headers)).coordinate}'

    for column_cells in worksheet.columns:
        max_length = 0
        column_letter = column_cells[0].column_letter

        for cell in column_cells:
            try:
                cell_length = len(str(cell.value)) if cell.value else 0
                if cell_length > max_length:
                    max_length = cell_length
            except Exception:
                pass

        worksheet.column_dimensions[column_letter].width = min(max_length + 4, 35)

    output = io.BytesIO()
    workbook.save(output)
    output.seek(0)

    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename="student_import_template_extended.xlsx"'

    return response

@never_cache
@students_import_export_access
def import_students_advanced(request):
    """استيراد متقدم للطلاب بنظام المعاينة قبل الحفظ"""

    if request.method == 'POST':
        action = request.POST.get('action', 'preview')

        # =========================
        # 1) معاينة الملف
        # =========================
        if action == 'preview':
            if 'file' not in request.FILES:
                messages.error(request, 'يرجى اختيار ملف للاستيراد')
                return redirect('students:import_students_advanced')

            file_obj = request.FILES['file']

            try:
                service = StudentImportPreviewService()
                summary = service.preview_file(file_obj)

                # حفظ الصفوف الصالحة في السيشن لحين تأكيد المستخدم
                request.session['student_import_valid_rows'] = summary['valid_rows']
                request.session['student_import_summary'] = {
                    'processed_count': summary['processed_count'],
                    'valid_count': summary['valid_count'],
                    'error_count': summary['error_count'],
                    'warning_count': summary['warning_count'],
                    'errors': summary['errors'],
                    'warnings': summary['warnings'],
                }
                request.session.modified = True

                context = {
                    'title': 'معاينة استيراد الطلاب',
                    'summary': summary,
                }

                return render(request, 'students/import_preview.html', context)

            except Exception as e:
                messages.error(request, f'حدث خطأ أثناء قراءة الملف: {str(e)}')
                return redirect('students:import_students_advanced')

        # =========================
        # 2) تأكيد الحفظ
        # =========================
        elif action == 'confirm':
            valid_rows = request.session.get('student_import_valid_rows', [])

            if not valid_rows:
                messages.error(request, 'لا توجد بيانات صالحة للحفظ. يرجى رفع الملف مرة أخرى.')
                return redirect('students:import_students_advanced')

            try:
                service = StudentImportPreviewService()
                result = service.confirm_import(valid_rows)

                # تنظيف السيشن بعد الحفظ
                request.session.pop('student_import_valid_rows', None)
                request.session.pop('student_import_summary', None)
                request.session.modified = True

                if result['created_count'] > 0:
                    messages.success(
                        request,
                        f'تم استيراد {result["created_count"]} طالب بنجاح'
                    )

                if result['errors']:
                    for error in result['errors'][:10]:
                        messages.error(request, error)

                return redirect('students:student_list')

            except Exception as e:
                messages.error(request, f'حدث خطأ أثناء حفظ الطلاب: {str(e)}')
                return redirect('students:import_students_advanced')

        # =========================
        # 3) إلغاء العملية
        # =========================
        elif action == 'cancel':
            request.session.pop('student_import_valid_rows', None)
            request.session.pop('student_import_summary', None)
            request.session.modified = True

            messages.info(request, 'تم إلغاء عملية الاستيراد')
            return redirect('students:import_students_advanced')

    context = {
        'title': 'استيراد بيانات الطلاب',
        'grade_levels': GradeLevel.objects.filter(
            is_active=True
        ).select_related(
            'education_level'
        ).order_by(
            'education_level__order',
            'order',
            'name'
        ),
    }

    return render(request, 'students/import_advanced.html', context)

@never_cache
@students_sensitive_operation
def upgrade_students_wizard(request):
    """معالج ترقية الطلاب للعام الجديد"""
    
    # التحقق من توفر الأدوات
    if not UPGRADE_AVAILABLE or not StudentUpgradeManager:
        messages.error(request, 'أدوات ترقية الطلاب غير متوفرة حالياً')
        return redirect('students:student_list')
    
    try:
        upgrade_manager = StudentUpgradeManager()
    except Exception as e:
        messages.error(request, f'خطأ في تحميل أدوات الترقية: {str(e)}')
        return redirect('students:student_list')
    
    # التحقق من إمكانية الترقية
    try:
        can_upgrade, message = upgrade_manager.can_perform_upgrade()
    except Exception as e:
        can_upgrade = False
        message = f"خطأ في فحص إمكانية الترقية: {str(e)}"
    
    if request.method == 'POST':
        if not can_upgrade:
            messages.error(request, message)
            return redirect('students:upgrade_students_wizard')
        
        action = request.POST.get('action')
        
        if action == 'preview':
            # عرض المعاينة
            try:
                preview_data = upgrade_manager.get_upgrade_preview()
                context = {
                    'preview_data': preview_data,
                    'can_upgrade': can_upgrade,
                    'title': 'معاينة ترقية الطلاب'
                }
                return render(request, 'students/upgrade_preview.html', context)
            except Exception as e:
                messages.error(request, f'خطأ في معاينة الترقية: {str(e)}')
        
        elif action == 'confirm':
            # تنفيذ الترقية
            try:
                selected_grades = request.POST.getlist('selected_grades')
                upgrade_options = {
                    'reset_fees': request.POST.get('reset_fees') == 'on',
                    'archive_graduates': request.POST.get('archive_graduates') == 'on'
                }
                
                success = upgrade_manager.perform_upgrade(
                    [int(id) for id in selected_grades],
                    upgrade_options
                )
                
                summary = upgrade_manager.get_upgrade_summary()
                
                if success:
                    messages.success(request, 
                        f'تم ترقية {summary["upgraded_count"]} طالب، '
                        f'تخرج {summary["graduated_count"]} طالب، '
                        f'أُرشف {summary["archived_count"]} طالب')
                
                if summary['errors']:
                    for error in summary['errors']:
                        messages.error(request, error)
                
                return redirect('students:student_list')
            except Exception as e:
                messages.error(request, f'خطأ في تنفيذ الترقية: {str(e)}')
    
    # الحصول على العام الحالي
    try:
        current_year = SettingsAcademicYear.get_current_year()
    except:
        current_year = None
    
    context = {
        'can_upgrade': can_upgrade,
        'message': message,
        'current_year': current_year,
        'title': 'ترقية الطلاب للعام الجديد'
    }
    return render(request, 'students/upgrade_wizard.html', context)

@never_cache
@students_basic_access
def user_guide(request):
    """دليل استخدام النظام - متاح لجميع المستخدمين"""
    user_role = get_user_role(request.user)
    system_settings, current_year = get_system_data()
    
    # معلومات دور المستخدم للدليل
    role_info = {
        'SYSTEM_ADMIN': {
            'title': 'المدير العام',
            'permissions': ['جميع العمليات', 'العمليات الحساسة', 'التقارير المتقدمة', 'الأدوات الإدارية'],
            'color': 'danger'
        },
        'SCHOOL_MANAGER': {
            'title': 'مدير المدرسة',
            'permissions': ['إدارة الطلاب', 'التقارير العامة', 'البيانات المالية'],
            'color': 'primary'
        },
        'STUDENT_AFFAIRS': {
            'title': 'موظف شؤون الطلاب',
            'permissions': ['إضافة الطلاب', 'البحث والعرض'],
            'color': 'success'
        },
        'TEACHER': {
            'title': 'المعلم',
            'permissions': ['العرض فقط', 'البحث في الطلاب'],
            'color': 'info'
        }
    }
    
    current_role_info = role_info.get(user_role, {
        'title': 'مستخدم',
        'permissions': ['العرض الأساسي'],
        'color': 'secondary'
    })
    
    # إحصائيات سريعة للدليل
    stats = {
        'total_students': Student.objects.filter(is_active=True).count(),
        'education_levels': EducationLevel.objects.filter(is_active=True).count(),
        'grade_levels': GradeLevel.objects.filter(is_active=True).count(),
    }
    
    context = {
        'user_role': user_role,
        'current_role_info': current_role_info,
        'system_settings': system_settings,
        'current_year': current_year,
        'stats': stats,
        'title': 'دليل استخدام نظام إدارة الطلاب',
    }
    return render(request, 'students/user_guide.html', context)

# ============================================================
# 📄 خطاب إفادة قيد طالب Word
# ============================================================

def set_paragraph_rtl(paragraph):
    """ضبط اتجاه الفقرة من اليمين لليسار"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    p_pr.append(bidi)


def set_run_font(run, font_name='Arial', size=14, bold=False):
    """ضبط خط عربي داخل ملف Word"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.append(r_fonts)

    r_fonts.set(qn('w:ascii'), font_name)
    r_fonts.set(qn('w:hAnsi'), font_name)
    r_fonts.set(qn('w:cs'), font_name)


def add_rtl_paragraph(document, text='', size=14, bold=False, alignment='right', space_after=6):
    """إضافة فقرة عربية من اليمين لليسار"""
    paragraph = document.add_paragraph()

    if alignment == 'center':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'left':
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    p_pr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    p_pr.append(bidi)

    paragraph.paragraph_format.space_after = Pt(space_after)

    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)

    return paragraph


def add_student_info_row(table, label, value):
    """إضافة صف بيانات داخل جدول الخطاب"""
    row = table.add_row()
    row.cells[0].text = str(value or 'غير محدد')
    row.cells[1].text = str(label)

    for cell in row.cells:
        for paragraph in cell.paragraphs:
            set_paragraph_rtl(paragraph)

            for run in paragraph.runs:
                set_run_font(run, size=12)

    # تمييز خانة العنوان
    for paragraph in row.cells[1].paragraphs:
        for run in paragraph.runs:
            run.bold = True


@never_cache
@students_basic_access
def student_enrollment_statement_word(request, pk):
    """
    تصدير خطاب إفادة قيد طالب بصيغة Word
    """
    student = get_object_or_404(
        Student.objects.select_related(
            'grade_level__education_level',
            'academic_year',
        ),
        pk=pk
    )

    system_settings, current_academic_year = get_system_data()

    school_name = 'مدرسة المنار الخاصة للغات'
    if system_settings and getattr(system_settings, 'school_name', None):
        school_name = system_settings.school_name

    document = Document()

    # إعداد الصفحة
    section = document.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # الهيدر
    add_rtl_paragraph(
        document,
        school_name,
        size=18,
        bold=True,
        alignment='center',
        space_after=2
    )

    add_rtl_paragraph(
        document,
        'إدارة شؤون الطلاب',
        size=13,
        bold=True,
        alignment='center',
        space_after=14
    )

    add_rtl_paragraph(
        document,
        'خطاب إفادة قيد طالب',
        size=20,
        bold=True,
        alignment='center',
        space_after=20
    )

    # التاريخ
    today = timezone.now().date()
    add_rtl_paragraph(
        document,
        f'تحريرًا في: {today.strftime("%Y/%m/%d")}',
        size=12,
        bold=False,
        alignment='left',
        space_after=18
    )

    # نص الخطاب
    add_rtl_paragraph(
        document,
        'تشهد إدارة المدرسة بأن الطالب / الطالبة الموضح بياناته أدناه مقيد بالمدرسة خلال العام الدراسي الحالي، وذلك طبقًا للبيانات المسجلة بسجلات شؤون الطلاب.',
        size=14,
        bold=False,
        alignment='right',
        space_after=14
    )

    # جدول البيانات
    table = document.add_table(rows=0, cols=2)
    table.style = 'Table Grid'

    identity_value = student.national_number or ''
    if not identity_value and getattr(student, 'passport_number', None):
        identity_value = f'جواز سفر: {student.passport_number}'

    grade_name = student.grade_level.name if student.grade_level else 'غير محدد'
    education_level_name = (
        student.grade_level.education_level.name
        if student.grade_level and student.grade_level.education_level
        else 'غير محدد'
    )

    academic_year_name = (
        student.academic_year.name
        if student.academic_year
        else current_academic_year.name if current_academic_year else 'غير محدد'
    )

    gender_display = (
        'ذكر' if student.gender == 'M'
        else 'أنثى' if student.gender == 'F'
        else 'غير محدد'
    )

    add_student_info_row(table, 'اسم الطالب', student.name)
    add_student_info_row(table, 'هوية الطالب', identity_value)
    add_student_info_row(table, 'الجنسية', getattr(student, 'nationality', '') or 'غير محدد')
    add_student_info_row(table, 'الديانة', student.get_religion_display() if getattr(student, 'religion', '') else 'غير محدد')
    add_student_info_row(table, 'النوع', gender_display)
    add_student_info_row(table, 'تاريخ الميلاد', student.date_of_birth.strftime('%Y/%m/%d') if student.date_of_birth else 'غير محدد')
    add_student_info_row(table, 'المرحلة التعليمية', education_level_name)
    add_student_info_row(table, 'الصف الدراسي', grade_name)
    add_student_info_row(table, 'العام الدراسي', academic_year_name)
    add_student_info_row(table, 'حالة القيد', student.get_enrollment_status_display() if hasattr(student, 'get_enrollment_status_display') else 'غير محدد')
    add_student_info_row(table, 'اسم ولي الأمر', student.parent_name or 'غير محدد')

    document.add_paragraph()

    # صيغة التأكيد
    add_rtl_paragraph(
        document,
        'وقد أعطيت هذه الإفادة بناءً على طلب ولي الأمر لتقديمها إلى من يهمه الأمر، دون أدنى مسؤولية على المدرسة بخلاف ما ورد بسجلاتها الرسمية.',
        size=14,
        bold=False,
        alignment='right',
        space_after=20
    )

    # توقيعات
    sign_table = document.add_table(rows=1, cols=3)
    sign_table.style = 'Table Grid'

    sign_table.cell(0, 0).text = 'ختم المدرسة'
    sign_table.cell(0, 1).text = 'مسؤول شؤون الطلاب'
    sign_table.cell(0, 2).text = 'مدير المدرسة'

    for row in sign_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_rtl(paragraph)
                for run in paragraph.runs:
                    set_run_font(run, size=12, bold=True)

    document.add_paragraph()
    add_rtl_paragraph(
        document,
        'يعتمد،',
        size=13,
        bold=True,
        alignment='left',
        space_after=4
    )

    # اسم الملف
    safe_student_name = ''.join(
        ch for ch in student.name
        if ch.isalnum() or ch in [' ', '_', '-']
    ).strip().replace(' ', '_')

    filename = f'enrollment_statement_{safe_student_name}_{student.pk}.docx'

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename}"'

    document.save(response)
    return response


def download_template(self, request):
    from openpyxl.worksheet.datavalidation import DataValidation
    from openpyxl.workbook.defined_name import DefinedName
    from openpyxl.utils import get_column_letter

    wb = openpyxl.Workbook()

    # =========================================
    # أدوات مساعدة
    # =========================================
    def sanitize_range_name(value):
        value = str(value or '').strip()
        if not value:
            return 'EMPTY'
        value = value.replace(' ', '_')
        value = value.replace('-', '_')
        value = value.replace('/', '_')
        value = value.replace('\\', '_')
        value = value.replace('(', '_')
        value = value.replace(')', '_')
        value = value.replace('.', '_')
        value = value.replace(',', '_')
        value = value.replace('،', '_')
        value = value.replace(':', '_')
        while '__' in value:
            value = value.replace('__', '_')
        if value and value[0].isdigit():
            value = f'LVL_{value}'
        return value

    header_fill = openpyxl.styles.PatternFill(
        start_color="4472C4",
        end_color="4472C4",
        fill_type="solid"
    )
    header_font = openpyxl.styles.Font(bold=True, color="FFFFFF")
    center_alignment = openpyxl.styles.Alignment(
        horizontal="center",
        vertical="center",
        wrap_text=True
    )

    # =========================================
    # جلب البيانات
    # =========================================
    academic_years = list(
        SettingsAcademicYear.objects.filter(is_active=True).order_by('-start_date', '-id')
    )

    grade_levels = list(
        GradeLevel.objects.filter(is_active=True)
        .select_related('education_level')
        .order_by('education_level__order', 'order', 'name')
    )

    education_levels = []
    education_level_map = {}
    seen_education_levels = set()

    for grade_level in grade_levels:
        education_name = grade_level.education_level.name if getattr(grade_level, 'education_level', None) else ''
        if not education_name:
            continue

        if education_name not in seen_education_levels:
            seen_education_levels.add(education_name)
            education_levels.append(education_name)

        education_level_map.setdefault(education_name, [])
        education_level_map[education_name].append(grade_level.name)

    # =========================================
    # الشيت 1: قالب الاستيراد
    # =========================================
    ws = wb.active
    ws.title = "قالب الاستيراد"
    ws.sheet_view.rightToLeft = True

    headers = self.get_import_headers()

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center_alignment

    sample_rows = [
        [
            'أحمد محمد علي',
            'طالب عادي',
            '30012010012345',
            '',
            'مصري',
            'مسلم',
            '14',
            'M',
            '2010-01-01',
            '01234567890',
            'القاهرة',
            academic_years[0].name if academic_years else '',
            education_levels[0] if education_levels else '',
            education_level_map.get(education_levels[0], [''])[0] if education_levels else '',
            'مستجد',
            '',
            '',
            'لا',
            '',
            'لا',
            'لا',
            'لا',
            '',
            'لا',
            '',
            '',
            'محمد علي أحمد',
            '01098765432',
            'parent@email.com',
            'محاسب',
            'الأب',
            '',
            '',
            '15000',
            '5000',
        ]
    ]

    for row_num, row_data in enumerate(sample_rows, 2):
        for col, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_num, column=col, value=value)
            cell.alignment = center_alignment

    instructions = [
        "تعليمات سريعة:",
        "1. استخدم نفس أسماء الأعمدة بدون تغيير.",
        "2. استخدم القوائم المنسدلة لاختيار العام الدراسي والمرحلة التعليمية والصف الدراسي.",
        "3. قائمة الصف الدراسي تعتمد على المرحلة التعليمية المختارة في نفس الصف.",
        "4. الاسم هو الحقل الإجباري الوحيد.",
        "5. الرقم القومي أو جواز السفر يجب ألا يكون مكرراً إذا تم إدخاله.",
        "6. النوع يقبل: M / F أو ذكر / أنثى.",
        "7. تاريخ الميلاد يفضل بصيغة YYYY-MM-DD.",
    ]

    start_row = len(sample_rows) + 4
    for index, instruction in enumerate(instructions):
        cell = ws.cell(row=start_row + index, column=1, value=instruction)
        if index == 0:
            cell.font = openpyxl.styles.Font(bold=True, color="D32F2F")
        else:
            cell.font = openpyxl.styles.Font(color="666666")

    ws.freeze_panes = 'A2'
    ws.auto_filter.ref = f'A1:{ws.cell(row=1, column=len(headers)).coordinate}'

    # =========================================
    # الشيت 2: الأعوام الدراسية
    # =========================================
    ws_years = wb.create_sheet(title="الاعوام الدراسية")
    ws_years.sheet_view.rightToLeft = True

    year_headers = ['م', 'العام الدراسي', 'نشط؟']
    for col, header in enumerate(year_headers, 1):
        cell = ws_years.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_alignment

    for row_num, academic_year in enumerate(academic_years, 2):
        ws_years.cell(row=row_num, column=1, value=row_num - 1)
        ws_years.cell(row=row_num, column=2, value=getattr(academic_year, 'name', str(academic_year)))
        ws_years.cell(row=row_num, column=3, value='نعم' if getattr(academic_year, 'is_active', False) else 'لا')

    ws_years.freeze_panes = 'A2'
    ws_years.auto_filter.ref = f'A1:C{max(ws_years.max_row, 2)}'

    # =========================================
    # الشيت 3: المراحل والصفوف
    # =========================================
    ws_grades = wb.create_sheet(title="المراحل والصفوف")
    ws_grades.sheet_view.rightToLeft = True

    grade_headers = ['م', 'المرحلة التعليمية', 'اسم النطاق', 'الصف الدراسي', 'نشط؟']
    for col, header in enumerate(grade_headers, 1):
        cell = ws_grades.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_alignment

    grade_row = 2
    for education_name, grades in education_level_map.items():
        range_name = sanitize_range_name(education_name)
        for idx, grade_name in enumerate(grades, 1):
            ws_grades.cell(row=grade_row, column=1, value=grade_row - 1)
            ws_grades.cell(row=grade_row, column=2, value=education_name)
            ws_grades.cell(row=grade_row, column=3, value=range_name)
            ws_grades.cell(row=grade_row, column=4, value=grade_name)
            ws_grades.cell(row=grade_row, column=5, value='نعم')
            grade_row += 1

    ws_grades.freeze_panes = 'A2'
    ws_grades.auto_filter.ref = f'A1:E{max(ws_grades.max_row, 2)}'

    # =========================================
    # الشيت 4: القيم المسموح بها
    # =========================================
    ws_choices = wb.create_sheet(title="قيم مسموح بها")
    ws_choices.sheet_view.rightToLeft = True

    choice_headers = ['نوع البيان', 'القيمة المعروضة', 'القيمة الداخلية']
    for col, header in enumerate(choice_headers, 1):
        cell = ws_choices.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = center_alignment

    row_num = 2

    for value, label in getattr(Student, 'STUDENT_TYPE_CHOICES', []):
        ws_choices.cell(row=row_num, column=1, value='نوع الطالب')
        ws_choices.cell(row=row_num, column=2, value=str(label))
        ws_choices.cell(row=row_num, column=3, value=str(value))
        row_num += 1

    for value, label in getattr(Student, 'RELIGION_CHOICES', []):
        ws_choices.cell(row=row_num, column=1, value='الديانة')
        ws_choices.cell(row=row_num, column=2, value=str(label))
        ws_choices.cell(row=row_num, column=3, value=str(value))
        row_num += 1

    for value, label in getattr(Student, 'ENROLLMENT_STATUS_CHOICES', []):
        ws_choices.cell(row=row_num, column=1, value='حالة القيد')
        ws_choices.cell(row=row_num, column=2, value=str(label))
        ws_choices.cell(row=row_num, column=3, value=str(value))
        row_num += 1

    for value, label in getattr(Student, 'EDUCATIONAL_GUARDIAN_CHOICES', []):
        ws_choices.cell(row=row_num, column=1, value='صاحب الولاية التعليمية')
        ws_choices.cell(row=row_num, column=2, value=str(label))
        ws_choices.cell(row=row_num, column=3, value=str(value))
        row_num += 1

    extra_values = [
        ('النوع', 'M', 'M'),
        ('النوع', 'F', 'F'),
        ('النوع', 'ذكر', 'M'),
        ('النوع', 'أنثى', 'F'),
        ('القيم المنطقية', 'نعم', 'True'),
        ('القيم المنطقية', 'لا', 'False'),
        ('القيم المنطقية', '1', 'True'),
        ('القيم المنطقية', '0', 'False'),
    ]

    for item_type, display_value, internal_value in extra_values:
        ws_choices.cell(row=row_num, column=1, value=item_type)
        ws_choices.cell(row=row_num, column=2, value=display_value)
        ws_choices.cell(row=row_num, column=3, value=internal_value)
        row_num += 1

    ws_choices.freeze_panes = 'A2'
    ws_choices.auto_filter.ref = f'A1:C{max(ws_choices.max_row, 2)}'

    # =========================================
    # شيت مخفي للقوائم
    # =========================================
    ws_lists = wb.create_sheet(title="_lists")
    ws_lists.sheet_view.rightToLeft = True

    student_type_labels = [str(label) for value, label in getattr(Student, 'STUDENT_TYPE_CHOICES', [])]
    religion_labels = [str(label) for value, label in getattr(Student, 'RELIGION_CHOICES', [])]
    enrollment_labels = [str(label) for value, label in getattr(Student, 'ENROLLMENT_STATUS_CHOICES', [])]
    guardian_labels = [str(label) for value, label in getattr(Student, 'EDUCATIONAL_GUARDIAN_CHOICES', [])]
    gender_values = ['M', 'F', 'ذكر', 'أنثى']
    yes_no_values = ['نعم', 'لا', '1', '0']

    for idx, value in enumerate(student_type_labels, 1):
        ws_lists.cell(row=idx, column=1, value=value)

    for idx, value in enumerate(religion_labels, 1):
        ws_lists.cell(row=idx, column=2, value=value)

    for idx, value in enumerate(enrollment_labels, 1):
        ws_lists.cell(row=idx, column=3, value=value)

    for idx, value in enumerate(guardian_labels, 1):
        ws_lists.cell(row=idx, column=4, value=value)

    for idx, value in enumerate(gender_values, 1):
        ws_lists.cell(row=idx, column=5, value=value)

    for idx, value in enumerate(yes_no_values, 1):
        ws_lists.cell(row=idx, column=6, value=value)

    for idx, academic_year in enumerate(academic_years, 1):
        ws_lists.cell(row=idx, column=7, value=getattr(academic_year, 'name', str(academic_year)))

    for idx, value in enumerate(education_levels, 1):
        ws_lists.cell(row=idx, column=8, value=value)

    max_grade_list_len = 1
    dependent_start_col = 9

    for offset, education_name in enumerate(education_levels):
        col_num = dependent_start_col + offset
        col_letter = get_column_letter(col_num)
        range_name = sanitize_range_name(education_name)

        ws_lists.cell(row=1, column=col_num, value=education_name)
        grade_names = education_level_map.get(education_name, [])
        max_grade_list_len = max(max_grade_list_len, len(grade_names))

        for grade_idx, grade_name in enumerate(grade_names, 2):
            ws_lists.cell(row=grade_idx, column=col_num, value=grade_name)

        end_row = max(len(grade_names) + 1, 2)
        wb.defined_names.add(
            DefinedName(
                range_name,
                attr_text=f"'_lists'!${col_letter}$2:${col_letter}${end_row}"
            )
        )

    wb.defined_names.add(
        DefinedName(
            'student_type_choices',
            attr_text=f"'_lists'!$A$1:$A${max(len(student_type_labels), 1)}"
        )
    )
    wb.defined_names.add(
        DefinedName(
            'religion_choices',
            attr_text=f"'_lists'!$B$1:$B${max(len(religion_labels), 1)}"
        )
    )
    wb.defined_names.add(
        DefinedName(
            'enrollment_choices',
            attr_text=f"'_lists'!$C$1:$C${max(len(enrollment_labels), 1)}"
        )
    )
    wb.defined_names.add(
        DefinedName(
            'guardian_choices',
            attr_text=f"'_lists'!$D$1:$D${max(len(guardian_labels), 1)}"
        )
    )
    wb.defined_names.add(
        DefinedName(
            'gender_choices',
            attr_text=f"'_lists'!$E$1:$E${max(len(gender_values), 1)}"
        )
    )
    wb.defined_names.add(
        DefinedName(
            'boolean_choices',
            attr_text=f"'_lists'!$F$1:$F${max(len(yes_no_values), 1)}"
        )
    )
    wb.defined_names.add(
        DefinedName(
            'academic_year_choices',
            attr_text=f"'_lists'!$G$1:$G${max(len(academic_years), 1)}"
        )
    )
    wb.defined_names.add(
        DefinedName(
            'education_level_choices',
            attr_text=f"'_lists'!$H$1:$H${max(len(education_levels), 1)}"
        )
    )

    ws_lists.sheet_state = 'hidden'

    # =========================================
    # Data Validation
    # =========================================
    max_rows = 5000

    dv_student_type = DataValidation(type="list", formula1="=student_type_choices", allow_blank=True)
    dv_student_type.prompt = "اختر نوع الطالب من القائمة"
    dv_student_type.error = "اختر قيمة صحيحة لنوع الطالب"
    ws.add_data_validation(dv_student_type)
    dv_student_type.add(f'B2:B{max_rows}')

    dv_religion = DataValidation(type="list", formula1="=religion_choices", allow_blank=True)
    dv_religion.prompt = "اختر الديانة من القائمة"
    dv_religion.error = "اختر قيمة صحيحة للديانة"
    ws.add_data_validation(dv_religion)
    dv_religion.add(f'F2:F{max_rows}')

    dv_gender = DataValidation(type="list", formula1="=gender_choices", allow_blank=True)
    dv_gender.prompt = "اختر النوع من القائمة"
    dv_gender.error = "اختر قيمة صحيحة للنوع"
    ws.add_data_validation(dv_gender)
    dv_gender.add(f'H2:H{max_rows}')

    dv_academic_year = DataValidation(type="list", formula1="=academic_year_choices", allow_blank=True)
    dv_academic_year.prompt = "اختر العام الدراسي من القائمة"
    dv_academic_year.error = "اختر عاماً دراسياً صحيحاً"
    ws.add_data_validation(dv_academic_year)
    dv_academic_year.add(f'L2:L{max_rows}')

    dv_education_level = DataValidation(type="list", formula1="=education_level_choices", allow_blank=True)
    dv_education_level.prompt = "اختر المرحلة التعليمية من القائمة"
    dv_education_level.error = "اختر مرحلة تعليمية صحيحة"
    ws.add_data_validation(dv_education_level)
    dv_education_level.add(f'M2:M{max_rows}')

    for row in range(2, max_rows + 1):
        dv_grade_level = DataValidation(
            type="list",
            formula1=f'=INDIRECT(SUBSTITUTE($M{row}," ","_"))',
            allow_blank=True
        )
        dv_grade_level.prompt = "اختر الصف الدراسي بناءً على المرحلة التعليمية"
        dv_grade_level.error = "اختر صفاً دراسياً من القائمة المرتبطة بالمرحلة التعليمية"
        ws.add_data_validation(dv_grade_level)
        dv_grade_level.add(f'N{row}')

    dv_enrollment = DataValidation(type="list", formula1="=enrollment_choices", allow_blank=True)
    dv_enrollment.prompt = "اختر حالة القيد من القائمة"
    dv_enrollment.error = "اختر حالة قيد صحيحة"
    ws.add_data_validation(dv_enrollment)
    dv_enrollment.add(f'O2:O{max_rows}')

    dv_boolean = DataValidation(type="list", formula1="=boolean_choices", allow_blank=True)
    dv_boolean.prompt = "اختر نعم أو لا"
    dv_boolean.error = "اختر قيمة صحيحة من القائمة"
    ws.add_data_validation(dv_boolean)
    for col in ['R', 'T', 'U', 'V', 'X']:
        dv_boolean.add(f'{col}2:{col}{max_rows}')

    dv_guardian = DataValidation(type="list", formula1="=guardian_choices", allow_blank=True)
    dv_guardian.prompt = "اختر صاحب الولاية التعليمية من القائمة"
    dv_guardian.error = "اختر قيمة صحيحة لصاحب الولاية التعليمية"
    ws.add_data_validation(dv_guardian)
    dv_guardian.add(f'AE2:AE{max_rows}')

    # =========================================
    # ضبط عرض الأعمدة
    # =========================================
    for sheet in wb.worksheets:
        for column in sheet.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    value_length = len(str(cell.value)) if cell.value else 0
                    if value_length > max_length:
                        max_length = value_length
                except Exception:
                    pass
            sheet.column_dimensions[column_letter].width = min(max_length + 4, 40)

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename="student_import_template.xlsx"'
    return response